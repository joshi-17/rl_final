from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("rel_final_report_updated.docx")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "FFFFFF")
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
    return p


def add_para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0)
        text = text[1:] if text.startswith(" ") else text
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    style_table(table)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)

    # Cover page
    for line, size, bold in [
        ("B. M. S. COLLEGE OF ENGINEERING", 16, True),
        ("Autonomous Institute, Affiliated to VTU, Belagavi", 12, False),
        ("DEPARTMENT OF MACHINE LEARNING", 14, True),
        ("Program: B.E. in Artificial Intelligence and Machine Learning", 12, False),
        ("Academic Year: 2025-26, Session: Jan 2026 - May 2026", 12, False),
        ("Reinforcement Learning (24AM6PCREL)", 12, True),
        ("Alternative Assessment Tool (AAT) Report", 14, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.bold = bold

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Smart Energy Management using Reinforcement Learning with MLOps Automation")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted by")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 0, 0)

    info = add_table(
        doc,
        ["Field", "Details"],
        [
            ["Student Name", ""],
            ["USN", ""],
            ["Course", "Reinforcement Learning (24AM6PCREL)"],
            ["Department", "Machine Learning"],
            ["Institution", "B. M. S. College of Engineering"],
        ],
    )
    add_page_break(doc)

    add_heading(doc, "Valuation Report", 1)
    add_para(doc, "(To be filled by faculty)")
    add_table(
        doc,
        ["Criterion", "Maximum Marks", "Marks Awarded", "Remarks"],
        [
            ["Data Preparation, Model Development and Evaluation", "6", "", ""],
            ["MLOps Automation and Scalability", "6", "", ""],
            ["Git and GitOps for Version Management and Collaboration", "5", "", ""],
            ["Documentation, Communication and Lifelong Learning", "3", "", ""],
        ],
    )
    add_page_break(doc)

    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Introduction",
        "2. Problem Statement",
        "3. Methodology",
        "4. Implementation Details",
        "5. MLOps, DevOps and Deployment Architecture",
        "6. Results and Interpretation",
        "7. Sustainable Development Goal Alignment",
        "8. Limitations and Future Scope",
        "9. Conclusion",
        "10. References",
        "Appendix A: Evidence Checklist",
    ]
    add_numbered(doc, toc)
    add_page_break(doc)

    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 Overview of Reinforcement Learning", 2)
    add_para(
        doc,
        "Reinforcement Learning is a machine learning approach in which an agent learns to make sequential decisions by interacting with an environment. The agent observes the current state, takes an action, receives a reward, and updates its behaviour so that cumulative reward improves over time.",
    )
    add_para(
        doc,
        "The project uses reinforcement learning for smart energy management in battery-powered IoT or edge devices. The device must decide how much power to allocate to incoming tasks while preserving battery life. This creates a natural trade-off between immediate task performance and long-term energy conservation.",
    )
    add_heading(doc, "1.2 Algorithm Used", 2)
    add_para(
        doc,
        "The implemented algorithm is tabular Q-learning. It was selected because the project has a compact discrete state space consisting of three battery levels and three task-demand levels. The complete state-action value table has dimensions 3 x 3 x 3, making the approach transparent, computationally efficient, and easy to interpret.",
    )
    add_para(doc, "The Q-learning update rule used in the implementation is:")
    add_para(doc, "Q(s, a) <- Q(s, a) + alpha [r + gamma max Q(s', a') - Q(s, a)]")

    add_heading(doc, "2. Problem Statement", 1)
    add_para(
        doc,
        "Modern battery-powered devices often execute tasks with varying computational demand. A fixed policy, such as always using medium power, wastes battery on low-demand tasks and may under-serve high-demand tasks. The objective of this project is to learn an adaptive policy that selects low, medium, or high power based on current battery state and task demand.",
    )
    add_heading(doc, "2.1 Expected Outcomes", 2)
    add_bullets(
        doc,
        [
            "Train a Q-learning agent that adapts power allocation to battery state and task demand.",
            "Compare the learned policy against a fixed medium-power baseline.",
            "Track experiments, metrics, model artifacts, and logs in a reproducible manner.",
            "Expose the trained policy through a FastAPI service for inference.",
            "Implement prediction logging and drift monitoring for deployed model behaviour.",
            "Provide CI/CD, Docker, DVC, and Kubernetes/GitOps assets for production-style delivery.",
        ],
    )

    add_heading(doc, "3. Methodology", 1)
    add_heading(doc, "3.1 Markov Decision Process Formulation", 2)
    add_table(
        doc,
        ["Element", "Description"],
        [
            ["State", "Tuple of battery level and task demand: (battery_level, task_demand)."],
            ["Battery Levels", "0 = low, 1 = medium, 2 = high."],
            ["Task Demand", "0 = low, 1 = medium, 2 = high."],
            ["Actions", "0 = low_power, 1 = medium_power, 2 = high_power."],
            ["Reward", "task_completion_score - battery_usage_penalty."],
            ["Terminal Condition", "Episode ends when battery reaches 0 percent."],
        ],
    )
    add_heading(doc, "3.2 Data Preparation and Feature Engineering", 2)
    add_para(
        doc,
        "The file data_prep.py implements a reproducible data-preparation stage. It generates simulated task records, removes invalid rows, clips battery values to the valid 0-100 range, and creates engineered features such as battery_bucket, task_demand_id, and is_high_demand. This makes the data pipeline explicit and auditable.",
    )
    add_heading(doc, "3.3 Training and Evaluation Process", 2)
    add_numbered(
        doc,
        [
            "Generate and clean task data using data_prep.py.",
            "Train Q-learning policies using YAML configuration files.",
            "Run hyperparameter tuning using tuning.py with multiple random seeds.",
            "Evaluate the trained policy against a fixed medium-power baseline.",
            "Generate comparison tables, plots, logs, and an HTML report.",
            "Register the trained policy and expose it through FastAPI.",
        ],
    )

    add_heading(doc, "4. Implementation Details", 1)
    add_table(
        doc,
        ["File or Folder", "Purpose"],
        [
            ["environment.py", "Defines the smart energy RL environment, state transitions, battery drain, and reward logic."],
            ["agent.py", "Implements the tabular Q-learning agent, epsilon-greedy action selection, Q-table update, save, and load logic."],
            ["baseline.py", "Implements the fixed baseline policy that always selects medium power."],
            ["train.py", "Runs config-driven training, logs results, writes per-episode metrics, tracks MLflow runs, and registers policies."],
            ["compare.py", "Evaluates RL policy against baseline and generates performance plots."],
            ["run_all.py", "Runs the full evaluation/reporting pipeline."],
            ["data_prep.py", "Performs data generation, cleaning, and feature engineering."],
            ["tuning.py", "Runs grid search and seed-based cross-validation for hyperparameter tuning."],
            ["api.py", "FastAPI service for health checks, predictions, and drift monitoring."],
            ["monitoring.py", "Logs prediction requests and computes task-distribution drift using KL divergence."],
            ["registry.py", "Maintains a simple local model registry with metadata for trained policies."],
            ["tests/", "Contains automated tests for agent behaviour, data preparation, drift detection, and API health."],
        ],
    )
    add_heading(doc, "4.1 Tools and Libraries", 2)
    add_table(
        doc,
        ["Tool", "Use in Project"],
        [
            ["Python", "Main implementation language."],
            ["NumPy", "Numerical operations and Q-table calculations."],
            ["Matplotlib", "Generation of reward and battery plots as output evidence."],
            ["PyYAML", "Configuration-driven experiments."],
            ["FastAPI", "REST API for deployed inference."],
            ["Uvicorn", "ASGI server for running FastAPI."],
            ["MLflow", "Experiment tracking for parameters, metrics, and artifacts."],
            ["Pytest", "Automated test execution."],
            ["Docker", "Containerized execution environment."],
            ["DVC", "Versioned ML pipeline stages."],
            ["Kubernetes", "Scalable API deployment and scheduled retraining manifests."],
        ],
    )

    add_heading(doc, "5. MLOps, DevOps and Deployment Architecture", 1)
    add_heading(doc, "5.1 Experiment Tracking and Model Registry", 2)
    add_para(
        doc,
        "The training script records every run in results.csv and stores detailed per-episode metrics in logs. It also logs hyperparameters, metrics, configuration files, and policy artifacts to MLflow. The registry.py module copies trained policy artifacts into a local model_registry folder with metadata, creating a lightweight model registry suitable for demonstration and audit.",
    )
    add_heading(doc, "5.2 CI/CD Pipeline", 2)
    add_para(
        doc,
        "The GitHub Actions workflow in .github/workflows/ci.yml implements continuous integration. On push or pull request, it checks out the code, installs dependencies, runs tests, performs a data-preparation smoke test, runs a training smoke test, starts the API, and verifies the /health endpoint. This prevents broken code from entering the main branch.",
    )
    add_heading(doc, "5.3 Docker and Docker Compose", 2)
    add_para(
        doc,
        "The Dockerfile defines a reproducible Python environment for the project. Docker Compose defines two services: a pipeline service for training/evaluation and an API service for FastAPI inference. This allows the same project to be run on another machine without manual environment setup.",
    )
    add_heading(doc, "5.4 FastAPI Serving Layer", 2)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["GET /health", "Checks whether the API is running and whether a policy artifact is available."],
            ["POST /predict", "Accepts battery percentage and task demand, then returns the selected power action."],
            ["GET /monitoring/drift", "Computes drift metrics from prediction logs and returns whether drift is detected."],
        ],
    )
    add_heading(doc, "5.5 Monitoring and Drift Detection", 2)
    add_para(
        doc,
        "Every API prediction is written to logs/predictions.csv. The monitoring endpoint compares live task-demand distribution against the expected training distribution using KL divergence. If the divergence exceeds the threshold, drift_detected is reported as true, indicating that retraining or model review may be required.",
    )
    add_heading(doc, "5.6 Kubernetes and GitOps", 2)
    add_para(
        doc,
        "The k8s folder contains Kubernetes manifests for ConfigMap, Deployment, Service, and CronJob. The Deployment runs the API with two replicas and uses /health for readiness and liveness probes. The Service provides stable network access. The CronJob represents scheduled retraining. Since these deployment files are stored in Git, they support a GitOps workflow where infrastructure changes are reviewed and version-controlled.",
    )
    add_heading(doc, "5.7 DVC Pipeline", 2)
    add_para(
        doc,
        "The dvc.yaml file defines four reproducible ML stages: prepare, tune, train, and evaluate. This captures the dependencies between source code, generated data, model artifacts, and reports. When an upstream dependency changes, DVC can rerun only the affected stages.",
    )

    add_heading(doc, "6. Results and Interpretation", 1)
    add_para(
        doc,
        "The RL policy is evaluated against a fixed medium-power baseline. The generated evidence is stored in outputs/index.html, outputs/comparison_table.md, and outputs/plots. Images are not inserted in this report as per the plain-report requirement, but the output files remain available in the project folder for examiner verification.",
    )
    add_table(
        doc,
        ["Metric", "Baseline", "RL Policy", "Interpretation"],
        [
            ["Average Reward", "Lower", "Higher", "The learned policy adapts actions to state and therefore improves cumulative reward."],
            ["Battery Remaining", "Usually 0 percent", "Usually 0 percent", "Episodes terminate on battery depletion; efficiency is better assessed through reward and drain per step."],
            ["Drain per Step", "Fixed medium drain", "Adaptive drain", "The RL agent can conserve power on low-demand tasks."],
            ["Task Adaptation", "No", "Yes", "The RL policy changes action depending on task demand."],
            ["Battery Awareness", "No", "Yes", "The RL policy includes battery level as part of the state."],
        ],
    )
    add_heading(doc, "6.1 Interpretation", 2)
    add_para(
        doc,
        "The baseline policy is simple but rigid. It always selects medium power, so it cannot exploit easy tasks or respond intelligently to high-demand tasks. The Q-learning policy learns from repeated interaction with the environment and gradually improves its action choices. Epsilon decay shifts the agent from exploration to exploitation as training progresses.",
    )
    add_para(
        doc,
        "The updated project also provides stronger reproducibility and observability than a basic RL implementation. Results can be reproduced through YAML configs, tracked using MLflow and CSV logs, served through FastAPI, monitored for drift, and deployed using Docker or Kubernetes manifests.",
    )

    add_heading(doc, "7. Sustainable Development Goal Alignment", 1)
    add_table(
        doc,
        ["SDG", "Relevance", "Project Contribution"],
        [
            ["SDG 7: Affordable and Clean Energy", "Energy efficiency", "Adaptive power allocation reduces unnecessary energy consumption in battery-powered devices."],
            ["SDG 11: Sustainable Cities and Communities", "Sustainable infrastructure", "Efficient IoT devices support scalable smart-city infrastructure with lower energy waste."],
            ["SDG 12: Responsible Consumption and Production", "Resource efficiency", "Longer battery life and fewer unnecessary charge cycles support responsible use of electronic resources."],
        ],
    )

    add_heading(doc, "8. Limitations and Future Scope", 1)
    add_bullets(
        doc,
        [
            "The environment is simulated and simplified; real hardware would introduce noise, thermal effects, and variable battery behaviour.",
            "Tabular Q-learning is suitable for small discrete state spaces but does not scale to high-dimensional continuous environments.",
            "The current monitoring detects distribution drift but does not automatically trigger retraining inside the API service.",
            "Future work can use DQN, PPO, or other deep reinforcement learning algorithms for richer environments.",
            "Future deployment can connect Kubernetes CronJob retraining to a real artifact store and production model registry.",
        ],
    )

    add_heading(doc, "9. Conclusion", 1)
    add_para(
        doc,
        "This project demonstrates a complete reinforcement learning and MLOps workflow for smart energy management. It begins with data preparation and feature engineering, trains and evaluates a Q-learning policy, compares it with a baseline, tracks experiments, registers model artifacts, exposes inference through FastAPI, monitors prediction drift, and includes CI/CD, Docker, DVC, and Kubernetes/GitOps assets. The result is a formal end-to-end project that goes beyond algorithm implementation and demonstrates how an RL model can be prepared for reproducible, scalable, and observable deployment.",
    )

    add_heading(doc, "10. References", 1)
    add_numbered(
        doc,
        [
            "Sutton, R. S., and Barto, A. G. Reinforcement Learning: An Introduction. MIT Press.",
            "Watkins, C. J. C. H., and Dayan, P. Q-learning. Machine Learning, 1992.",
            "FastAPI Documentation. https://fastapi.tiangolo.com/",
            "MLflow Documentation. https://mlflow.org/docs/latest/",
            "Docker Documentation. https://docs.docker.com/",
            "Kubernetes Documentation. https://kubernetes.io/docs/",
            "DVC Documentation. https://dvc.org/doc",
            "United Nations Sustainable Development Goals. https://sdgs.un.org/goals",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "Appendix A: Evidence Checklist", 1)
    add_table(
        doc,
        ["Rubric Requirement", "Implemented Evidence"],
        [
            ["Data cleaning and feature engineering", "data_prep.py, data/raw_tasks.csv, data/clean_tasks.csv"],
            ["Multiple models", "baseline.py and agent.py"],
            ["Hyperparameter tuning and cross-validation", "tuning.py, outputs/tuning_results.csv, outputs/best_hyperparameters.json"],
            ["Relevant metrics", "results.csv, logs/*.json, outputs/comparison_table.md"],
            ["MLflow tracking", "mlruns/ generated by train.py"],
            ["Reproducible script", "run_all.py and YAML configs"],
            ["CI/CD", ".github/workflows/ci.yml"],
            ["Docker orchestration", "Dockerfile and docker-compose.yml"],
            ["REST API", "api.py with /health, /predict, and /monitoring/drift"],
            ["Monitoring", "monitoring.py, logs/predictions.csv, logs/drift_metrics.jsonl"],
            ["GitOps and Kubernetes", "k8s/configmap.yaml, deployment.yaml, service.yaml, retrain-cronjob.yaml"],
            ["Versioning", "dvc.yaml and model_registry/"],
            ["Collaboration", ".github/PULL_REQUEST_TEMPLATE.md and .github/ISSUE_TEMPLATE.md"],
        ],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Smart Energy Management using Reinforcement Learning - AAT Report"
        for run in footer.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
